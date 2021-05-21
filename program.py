import argparse
import sys
from docx import Document
from docx.shared import RGBColor

class MyError(Exception):
    def __init__(self, text):
        self.txt = text
       
       
def start(orig, secret):
    secretStr = ''
    secretFile = open(secret, 'r')
    secretStr += secretFile.read()
    secretFile.close()
    secretSymbols = []
    
    for i in secretStr.replace(' ', ''):
        secretSymbols.append(i.lower())
    
    origDoc = Document(orig)
    docLength = getDocLength(origDoc)
   
    if docLength < len(secretSymbols):
        raise MyError('The secret message cannot be inserted into this document, because the message is longer than the text.')
    else:
        createSecretDoc(origDoc, secretSymbols)
        return

    
def getDocLength(doc):
    docLength = 0
    for paragraph in doc.paragraphs:
        parWithoutSpace = paragraph.text.replace(' ', '')
        docLength += len(parWithoutSpace)
    return docLength


def createSecretDoc(origDoc, secretSymbols):
    secretDoc = Document()
    for paragraph in origDoc.paragraphs:
        parOrigText = paragraph.text
        parSecretDoc = secretDoc.add_paragraph()
        for char in parOrigText:
            docSymbol = parSecretDoc.add_run(f'{char}')
            changeColor(secretSymbols, char, docSymbol)
    if len(secretSymbols) != 0: 
        raise MyError('The secret message cannot be inserted into this document.')
    secretDoc.save('secretText.docx')

    
def changeColor(secretSymbols, char, docSymbol):
    font = docSymbol.font
    if len(secretSymbols) != 0 and char.lower() == secretSymbols[0].lower():
        secretSymbols.pop(0)
        font.color.rgb = RGBColor(0xFF, 0x01, 0x01)
    else:
        font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    
def createParser():
    parser = argparse.ArgumentParser()
    parser.add_argument('text',
                        help='docx file')
    parser.add_argument('secret',
                        help='txt secret message')
    return parser

    
if __name__ == '__main__':
    namespase = createParser().parse_args(sys.argv[1:])
    start(namespase.text, namespase.secret)
